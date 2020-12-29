"""
features missing from the docx library
"""
from typing import Tuple

import docx.document
import docx.table
from docx.table import Table


class CachedTable(Table):
    """
    deal with bad table performance of docx library
    https://github.com/python-openxml/python-docx/issues/174#issuecomment-667784193

    usage:

    ROWS=1000
    COLUMNS = 7
    document=docx.Document()
    table = CachedTable.transform(document.add_table(rows=ROWS, columns=COLUMNS))
    for i in range(ROWS):
        for j in range(COLUMNS):
            cell = table.cell(i, j)
            # Add text to row_cells
    """

    def __init__(self, tbl, parent):
        """

        @param tbl:
        @param parent:
        """
        super().__init__(tbl, parent)
        self._element = self._tbl = tbl
        self._cached_cells = None

    @property
    def _cells(self):
        """

        @return:
        """
        if self._cached_cells is None:
            self._cached_cells = super()._cells
        return self._cached_cells

    @staticmethod
    def transform(table):
        """

        @param table:
        @return:
        """
        # noinspection PyProtectedMember
        cached_table = CachedTable(table._tbl, table._parent)  # pylint: disable=protected-access
        return cached_table


def add_report_table(document: docx.document.Document, headings: Tuple, n_rows: int) -> CachedTable:
    """
    insert empty table with heading
    @param document: insert point
    @param headings: tuple with heading of a table (controls number of columns)
    @param n_rows: number of rows
    @return:
    """
    table = CachedTable.transform(document.add_table(rows=n_rows, cols=len(headings), style="Game Schedule Table"))
    for i, heading in enumerate(headings):
        table.cell(0, i).text = heading
    return table
